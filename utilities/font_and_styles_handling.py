from collections import defaultdict
from docx.enum.text import WD_COLOR_INDEX
from utilities.utils import estimate_page_number , format_errors

def highlight_font_size_error(paragraph):
    for run in paragraph.runs:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW

def check_font_size(document):
    errors = defaultdict(lambda: defaultdict(list))
    standard_font_size_body = 12
    standard_font_size_heading_1 = 14

    for i, para in enumerate(document.paragraphs):
        expected_font_size = standard_font_size_heading_1 if para.style.name == 'Heading 1' else standard_font_size_body
        for run in para.runs:
            run_font_size = run.font.size.pt if run.font.size else expected_font_size
            if run_font_size != expected_font_size:
                highlight_font_size_error(para)
                page_num = estimate_page_number(i)
                errors[page_num][(expected_font_size, run_font_size)].append(i + 1)

    formatted_errors = ["Font Size Errors", f"Total Font Size Errors: {sum(len(paras) for page in errors.values() for paras in page.values())}"]

    for page_num, page_errors in sorted(errors.items()):
        for (expected, found), paragraphs in sorted(page_errors.items()):
            formatted_errors.append(f"Error: Incorrect font size (Expected: {expected} pt, Found: {found} pt)")
            formatted_errors.append(f"• Page {page_num}: Paragraphs {', '.join(map(str, sorted(set(paragraphs))))}") # Use set to remove duplicates
            formatted_errors.append("")  # Add a blank line for readability

    return formatted_errors

def check_font_style(document):
    errors = defaultdict(lambda: defaultdict(list))
    standard_font = "Arial"

    for i, para in enumerate(document.paragraphs):
        for run in para.runs:
            run_font = run.font.name if run.font.name else standard_font
            if run_font.lower() != standard_font.lower():
                page_num = estimate_page_number(i)
                errors[page_num][(standard_font, run_font)].append(i + 1)

    formatted_errors = ["Font Style Errors", f"Total Font Style Errors: {sum(len(paras) for page in errors.values() for paras in page.values())}"]

    for page_num, page_errors in sorted(errors.items()):
        for (expected, found), paragraphs in sorted(page_errors.items()):
            formatted_errors.append(f"Error: Incorrect font style (Expected: {expected}, Found: {found})")
            formatted_errors.append(f"• Page {page_num}: Paragraphs {', '.join(map(str, sorted(set(paragraphs))))}")
            formatted_errors.append("")  

    return formatted_errors 

def check_caption_styles(doc):
    errors = defaultdict(lambda: defaultdict(list))
    standard_font = "Arial"
    standard_size = 11 * 12700  # Word font size is in half-points (1 pt = 12700 EMUs)

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip().lower()
        
        # Identify table and figure captions
        if text.startswith("table") or text.startswith("figure"):
            for run in para.runs:
                run_font = run.font.name if run.font and run.font.name else standard_font
                run_size = run.font.size if run.font and run.font.size else standard_size
                
                page_num = estimate_page_number(i)

                # Check font style
                if run_font.lower() != standard_font.lower():
                    errors[page_num]["Font Style"].append(i + 1)

                # Check font size
                if run_size != standard_size:
                    errors[page_num]["Font Size"].append(i + 1)

    formatted_errors = ["Table & Figure Caption Errors"]
    total_errors = sum(len(paras) for page in errors.values() for paras in page.values())
    formatted_errors.append(f"Total Errors: {total_errors}\n")

    for page_num, page_errors in sorted(errors.items()):
        for error_type, paragraphs in sorted(page_errors.items()):
            formatted_errors.append(f"Error: Incorrect {error_type} (Expected: {standard_font}, 11pt)")
            formatted_errors.append(f"• Page {page_num}: Paragraphs {', '.join(map(str, sorted(set(paragraphs))))}\n")

    return formatted_errors



def check_heading_styles(doc):
    """Check if headings use the correct styles (Heading 1 for titles, Heading 2 for subheadings)."""
    errors = defaultdict(list)

    for i, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else "Normal"

        if para.text.strip():  # Ignore empty paragraphs
            if "heading" in style_name.lower():
                if style_name.lower() == "heading 1":
                    continue  # Correct style for titles
                elif style_name.lower() == "heading 2":
                    continue  # Correct style for subheadings
                else:
                    errors["Incorrect Heading Style"].append(i + 1)
            elif para.text.isupper():  # Titles should not be in Normal text
                errors["Title in Normal Style"].append(i + 1)

    return format_errors(errors, "Heading Style Errors")

